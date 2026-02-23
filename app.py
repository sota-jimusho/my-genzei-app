import streamlit as st
from docx import Document
import io

# アプリの設定
st.set_page_config(page_title="住宅ローン減税シミュレーター", layout="wide")
st.title("🏠 住宅ローン減税シミュレーター")

# --- 1. 入力セクション（サイドバー） ---
st.sidebar.header("【入力セクション】")
nenshu = st.sidebar.number_input("年収（万円）を入力してください", value=600, step=10)
loan_zandaka = st.sidebar.number_input("住宅ローンの年末残高（万円）", value=3500, step=100)

st.sidebar.markdown("---")
st.sidebar.header("【住宅ローン条件入力】")
setai_type = st.sidebar.radio("子育て・若者世帯ですか？", ["1: はい", "2: いいえ"])
jutaku_kubun = st.sidebar.radio("住宅は？", ["1: 新築", "2: 中古"])
seino_type = st.sidebar.selectbox("性能は？", ["1: 長期優良", "2: ZEH", "3: 省エネ", "4: その他"])

# --- 2. 計算ロジック（元のコードを忠実に再現） ---

# 給与所得控除の計算
if nenshu <= 180:
    kyuyo_kojo = nenshu * 0.4 - 10
    k_formula = f"{nenshu} * 0.4 - 10"
elif nenshu <= 360:
    kyuyo_kojo = nenshu * 0.3 + 8
    k_formula = f"{nenshu} * 0.3 + 8"
elif nenshu <= 660:
    kyuyo_kojo = nenshu * 0.2 + 44
    k_formula = f"{nenshu} * 0.2 + 44"
elif nenshu <= 850:
    kyuyo_kojo = nenshu * 0.1 + 110
    k_formula = f"{nenshu} * 0.1 + 110"
else:
    kyuyo_kojo = 195
    k_formula = "上限 195"

gokei_shotoku = nenshu - kyuyo_kojo

# 基礎控除の計算
if gokei_shotoku <= 132: kiso_kojo = 95
elif gokei_shotoku <= 336: kiso_kojo = 88
elif gokei_shotoku <= 489: kiso_kojo = 68
elif gokei_shotoku <= 655: kiso_kojo = 63
elif gokei_shotoku <= 2350: kiso_kojo = 58
else: kiso_kojo = 0

# 社会保険料（15%）
shakai_hoken = nenshu * 0.15
shotoku_kojo_gokei = kiso_kojo + shakai_hoken

# オプション控除
st.header("📋 その他の所得控除の設定")
c1, c2 = st.columns(2)
with c1:
    ideco_check = st.checkbox("iDeCoをやっていますか？")
    ideco_nenkan = 0
    if ideco_check:
        getzugaku = st.number_input("毎月の掛金（円）", value=23000)
        ideco_nenkan = (getzugaku * 12) / 10000
        shotoku_kojo_gokei += ideco_nenkan
with c2:
    fuyo_check = st.checkbox("16歳以上の扶養家族はいますか？")
    fuyo_kojo = 0
    if fuyo_check:
        ninzu = st.number_input("人数", min_value=1, value=1)
        fuyo_kojo = ninzu * 38
        shotoku_kojo_gokei += fuyo_kojo

# 課税所得と所得税（控除前）
kazei_shotoku = max(0, gokei_shotoku - shotoku_kojo_gokei)

if kazei_shotoku <= 195:
    shotokuzei_mae = kazei_shotoku * 0.05
    s_formula = f"{kazei_shotoku:.1f} * 0.05"
elif kazei_shotoku <= 330:
    shotokuzei_mae = kazei_shotoku * 0.1 - 9.75
    s_formula = f"{kazei_shotoku:.1f} * 0.1 - 9.75"
elif kazei_shotoku <= 695:
    shotokuzei_mae = kazei_shotoku * 0.2 - 42.75
    s_formula = f"{kazei_shotoku:.1f} * 0.2 - 42.75"
elif kazei_shotoku <= 900:
    shotokuzei_mae = kazei_shotoku * 0.23 - 63.6
    s_formula = f"{kazei_shotoku:.1f} * 0.23 - 63.6"
else:
    shotokuzei_mae = kazei_shotoku * 0.33 - 153.6
    s_formula = f"{kazei_shotoku:.1f} * 0.33 - 153.6"

# 住宅ローン控除限度額の判定
gendo_gaku = 2000 
s_type = "1" if setai_type == "1: はい" else "2"
j_type = "1" if jutaku_kubun == "1: 新築" else "2"
p_type = seino_type[0] # "1", "2", "3", "4"を取得

if j_type == "1":
    if p_type == "1": gendo_gaku = 5000 if s_type == "1" else 4500
    elif p_type == "2": gendo_gaku = 4500 if s_type == "1" else 3500
    elif p_type == "3": gendo_gaku = 3000 if s_type == "1" else 2000
elif j_type == "2":
    if p_type in ["1", "2"]: gendo_gaku = 3000
    else: gendo_gaku = 2000

# 控除適用
loan_kojo_waku = min(loan_zandaka, gendo_gaku) * 0.007
actual_shotoku_deduction = min(shotokuzei_mae, loan_kojo_waku)
remaining_waku = loan_kojo_waku - actual_shotoku_deduction
juminzei_genkai = min(kazei_shotoku * 0.05, 9.75) 
juminzei_mae = (kazei_shotoku * 0.1) + 0.5 
actual_jumin_deduction = min(remaining_waku, juminzei_genkai, juminzei_mae)

# 最大借入額の逆算
tax_max_benefit = shotokuzei_mae + juminzei_genkai
house_max_benefit = gendo_gaku * 0.007
real_max_benefit = min(tax_max_benefit, house_max_benefit)
optimal_loan_amount = real_max_benefit / 0.007

# --- 3. 結果表示セクション（元の文章を使用） ---
st.divider()
st.header("■ 計算結果")

res_col1, res_col2 = st.columns(2)
with res_col1:
    st.write(f"合計所得金額: {gokei_shotoku:.1f}万円")
    st.write(f"課税所得金額: {kazei_shotoku:.1f}万円※社会保険料を年収の15％とした概算")
    st.write(f"住宅ローン減税合計額: {actual_shotoku_deduction + actual_jumin_deduction:.1f}万円")
    st.write(f"所得税金額: {shotokuzei_mae:.2f}万円")
    st.write(f"住宅ローン控除額（所得税）: {actual_shotoku_deduction:.2f}万円")
    st.write(f"控除後所得税額: {max(0, shotokuzei_mae - actual_shotoku_deduction):.2f}万円")
with res_col2:
    st.write(f"住民税金額: {juminzei_mae:.2f}万円")
    st.write(f"住宅ローン控除額（住民税）: {actual_jumin_deduction:.2f}万円")
    st.write(f"控除後住民税額: {max(0, juminzei_mae - actual_jumin_deduction):.2f}万円")
    st.write(f"控除余剰額 : {max(0, loan_kojo_waku - (actual_shotoku_deduction + actual_jumin_deduction)):.2f}万円")

st.success(f"★ 控除を最大限活かせる借入額の目安: 約 {int(optimal_loan_amount)} 万円")
if optimal_loan_amount > gendo_gaku:
    st.warning(f"※注: あなたの納税額に対して住宅の限度額({gendo_gaku}万)が上限です。")

# --- 4. 計算プロセスの明細（元の文章を使用） ---
with st.expander("【 実際の計算プロセス明細を確認する 】"):
    st.write(f"1. 所得算出: {nenshu}（年収） - ({k_formula}（給与所得控除）) = {gokei_shotoku:.1f}万円")
    st.write(f"2. 課税所得: {gokei_shotoku:.1f} - {kiso_kojo}(基礎控除) - {shakai_hoken:.1f}(社保控除) - {ideco_nenkan + fuyo_kojo:.1f}(他控除) = {kazei_shotoku:.1f}万円")
    st.write(f"3. 所得税額: {s_formula} = {shotokuzei_mae:.2f}万円")
    st.write(f"4. ローン控除内訳:")
    st.write(f"   - 合計額: ({loan_zandaka}万（年末ローン残高）, {gendo_gaku}万（住宅種別による控除限度額）のいずれか少ない方) * 0.7% = {loan_kojo_waku:.2f}万円")
    st.write(f"   - 所得税から: {actual_shotoku_deduction:.2f}万円")
    st.write(f"   - 住民税から: {actual_jumin_deduction:.2f}万円 (ローン控除合計額-控除額（所得税）:{remaining_waku:.2f}/上限（9.75万円又は課税所得の0.5％いずれか少ない方）:{juminzei_genkai:.2f}/住民税額:{juminzei_mae:.2f} の最小値)")
    st.write(f"5. 逆算根拠: (所得税{shotokuzei_mae:.2f}万 + 住民税上限{juminzei_genkai:.2f}万) / 0.7% = {optimal_loan_amount:.1f}万円")

# --- 5. Wordダウンロード ---
st.write("---")
if st.button("simulation_result.docx を作成してダウンロード"):
    doc = Document()
    doc.add_heading('住宅ローン減税計算結果', 0)
    doc.add_heading('■ 計算結果', level=1)
    res_text = (
        f"合計所得金額: {gokei_shotoku:.1f}万円\n"
        f"課税所得金額: {kazei_shotoku:.1f}万円\n"
        f"所得税金額: {shotokuzei_mae:.2f}万円\n"
        f"住宅ローン控除額（所得税）: {actual_shotoku_deduction:.2f}万円\n"
        f"控除後所得税額: {max(0, shotokuzei_mae - actual_shotoku_deduction):.2f}万円\n"
        f"住民税金額: {juminzei_mae:.2f}万円\n"
        f"住宅ローン控除額（住民税）: {actual_jumin_deduction:.2f}万円\n"
        f"最終 住民税額: {max(0, juminzei_mae - actual_jumin_deduction):.2f}万円\n"
        f"控除余剰額   : {max(0, loan_kojo_waku - (actual_shotoku_deduction + actual_jumin_deduction)):.2f}万円\n"
        f"----------------------------------------\n"
        f"★ 控除を最大限活かせる借入額の目安: 約 {int(optimal_loan_amount)} 万円\n"
    )
    doc.add_paragraph(res_text)
    
    doc.add_heading('■ 計算式明細', level=1)
    detail_text = (
        f"1. 所得算出: {nenshu}（年収） - ({k_formula}（給与所得控除）) = {gokei_shotoku:.1f}万円\n"
        f"2. 課税所得: {gokei_shotoku:.1f} - {kiso_kojo}(基礎控除) - {shakai_hoken:.1f}(社保控除) - {ideco_nenkan + fuyo_kojo:.1f}(その他控除) = {kazei_shotoku:.1f}万円\n"
        f"3. 所得税額: {s_formula} = {shotokuzei_mae:.2f}万円\n"
        f"4. ローン控除内訳:\n"
        f"   - 合計額: ({loan_zandaka}万, {gendo_gaku}万のいずれか少ない方) * 0.7% = {loan_kojo_waku:.2f}万円\n"
        f"   - 所得税から: {actual_shotoku_deduction:.2f}万円\n"
        f"   - 住民税から: {actual_jumin_deduction:.2f}万円\n"
    )
    doc.add_paragraph(detail_text)

    bio = io.BytesIO()
    doc.save(bio)

    st.download_button(label="📥 Wordファイルをダウンロード", data=bio.getvalue(), file_name="simulation_result.docx")


